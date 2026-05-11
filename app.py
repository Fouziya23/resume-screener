from flask import Flask, request, render_template
from PyPDF2 import PdfReader
from docx import Document
from rapidfuzz import fuzz
import re, math

app = Flask(__name__)

# ── Stop words ────────────────────────────────────────────────────────────────
STOP_WORDS = set([
    "a","an","the","and","or","but","in","on","at","to","for","of","with",
    "by","from","is","are","was","were","be","been","being","have","has",
    "had","do","does","did","will","would","could","should","may","might",
    "shall","can","need","must","am","it","its","this","that","these","those",
    "we","our","you","your","they","their","i","my","he","she","his","her",
    "as","if","so","not","no","nor","yet","both","either","each","more",
    "most","other","some","such","than","then","when","where","who","which",
    "how","all","any","few","into","through","during","before","after",
    "above","below","between","out","off","over","under","again","further",
    "there","here","just","also","about","up","down","what","only","same",
    "use","used","using","work","working","including","etc","e.g","i.e"
])

# ── Skill taxonomy with categories ───────────────────────────────────────────
SKILL_TAXONOMY = {
    "Database & SQL": {
        "sql": ["sql", "mysql", "postgresql", "sql server", "mssql", "t-sql", "pl/sql", "sqlite", "oracle db"],
        "nosql": ["nosql", "mongodb", "cassandra", "dynamodb", "redis"],
        "database": ["database", "db", "rdbms", "data warehouse", "snowflake", "bigquery", "redshift"],
    },
    "Programming": {
        "python": ["python", "pandas", "numpy", "matplotlib", "seaborn", "scikit-learn", "scipy", "jupyter"],
        "r": ["r programming", "ggplot", "dplyr", "tidyr", "r studio"],
        "other": ["java", "scala", "c++", "javascript", "vba", "macro"],
    },
    "BI & Visualization": {
        "power bi": ["power bi", "powerbi", "dax", "power query", "power pivot"],
        "tableau": ["tableau", "tableau desktop", "tableau server"],
        "excel": ["excel", "vlookup", "xlookup", "index match", "pivot table", "power query", "advanced excel"],
        "other bi": ["looker", "qlik", "google data studio", "metabase", "superset"],
    },
    "Analytics & Modeling": {
        "machine learning": ["machine learning", "ml", "random forest", "decision tree", "xgboost", "classification", "regression", "clustering", "neural network"],
        "statistics": ["statistics", "statistical analysis", "hypothesis testing", "p-value", "confidence interval", "anova", "regression analysis"],
        "analytics": ["eda", "exploratory data analysis", "predictive modeling", "forecasting", "cohort analysis", "rfm", "a/b testing", "churn analysis"],
    },
    "Data Engineering": {
        "etl": ["etl", "data pipeline", "data ingestion", "airflow", "spark", "pyspark", "hadoop", "kafka"],
        "cloud": ["aws", "azure", "gcp", "google cloud", "s3", "ec2", "databricks"],
    },
    "Soft & Business Skills": {
        "communication": ["communication", "stakeholder", "presentation", "storytelling", "report writing"],
        "business": ["business analysis", "requirements gathering", "kpi", "dashboard", "mis reporting", "data driven"],
        "tools": ["git", "github", "jira", "confluence", "agile", "scrum"],
    }
}

# Flatten for matching
def get_all_aliases():
    result = {}
    for category, skills in SKILL_TAXONOMY.items():
        for canonical, aliases in skills.items():
            result[canonical] = {"aliases": aliases, "category": category}
    return result

ALL_SKILLS = get_all_aliases()

# Stemming rules (simple suffix stripping)
def simple_stem(word):
    word = word.lower()
    for suffix in ["ing", "tion", "ations", "ized", "ises", "ical", "ively", "ness", "ment", "ed", "er", "ly", "al", "s"]:
        if word.endswith(suffix) and len(word) - len(suffix) > 3:
            return word[:-len(suffix)]
    return word

# ── Text extraction ───────────────────────────────────────────────────────────
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx_file(file):
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text(file):
    name = file.filename.lower()
    file.seek(0)
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif name.endswith(".docx"):
        return extract_text_from_docx_file(file)
    return ""

# ── Keyword extraction from JD ────────────────────────────────────────────────
def extract_jd_keywords(text):
    text_lower = text.lower()

    # Extract skill phrases
    skill_matches = {}
    for canonical, info in ALL_SKILLS.items():
        for alias in info["aliases"]:
            if alias in text_lower:
                # count frequency
                count = len(re.findall(re.escape(alias), text_lower))
                if canonical not in skill_matches or skill_matches[canonical]["count"] < count:
                    skill_matches[canonical] = {
                        "count": count,
                        "category": info["category"],
                        "matched_alias": alias
                    }

    # Extract other important words (nouns/roles)
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9\+\#\.]*\b', text)
    freq = {}
    for w in words:
        wl = w.lower()
        if wl not in STOP_WORDS and len(wl) > 3:
            freq[wl] = freq.get(wl, 0) + 1

    skip = {"experience","work","team","role","company","job","position","candidate","required",
            "preferred","including","responsibilities","qualifications","requirements","ability",
            "strong","good","knowledge","understanding","working","based","across","within",
            "environment","opportunity","looking","hire","please","apply","must","will","also",
            "well","both","their","other","year","years","month","day","time","level","high"}

    top_words = [(k, v) for k, v in sorted(freq.items(), key=lambda x: -x[1])
                 if k not in skip and k not in [c for c in skill_matches]][:25]

    return skill_matches, top_words

# ── Resume skill matching ─────────────────────────────────────────────────────
def match_resume_skills(resume_text, skill_matches):
    resume_lower = resume_text.lower()
    resume_words = set(simple_stem(w) for w in re.findall(r'\b\w+\b', resume_lower))

    matched = {}
    missing = {}

    for canonical, info in skill_matches.items():
        found = False
        # 1. Direct alias match
        for alias in ALL_SKILLS[canonical]["aliases"]:
            if alias in resume_lower:
                found = True
                break
        # 2. Fuzzy match on canonical
        if not found:
            score = fuzz.partial_ratio(canonical, resume_lower)
            if score >= 88:
                found = True
        # 3. Stem match
        if not found:
            canonical_stem = simple_stem(canonical)
            if canonical_stem in resume_words:
                found = True

        if found:
            matched[canonical] = info
        else:
            missing[canonical] = info

    return matched, missing

def match_general_keywords(resume_text, top_words):
    resume_lower = resume_text.lower()
    resume_stems = set(simple_stem(w) for w in re.findall(r'\b\w+\b', resume_lower))
    matched_words, missing_words = [], []
    for word, freq in top_words:
        if word in resume_lower or simple_stem(word) in resume_stems:
            matched_words.append(word)
        else:
            missing_words.append(word)
    return matched_words, missing_words

# ── Section detection ─────────────────────────────────────────────────────────
SECTION_PATTERNS = {
    "Summary/Objective": [r"summary", r"objective", r"profile", r"about me"],
    "Work Experience": [r"experience", r"work experience", r"employment", r"internship"],
    "Education": [r"education", r"academic", r"qualification", r"degree"],
    "Skills": [r"skills", r"technical skills", r"core skills", r"competencies"],
    "Projects": [r"projects", r"portfolio", r"case studies"],
    "Certifications": [r"certif", r"courses", r"training", r"accreditation"],
    "Contact Info": [r"@", r"linkedin", r"phone", r"\d{10}"],
    "Achievements": [r"achievement", r"award", r"honor", r"recognition"],
}

def detect_sections(resume_text):
    text_lower = resume_text.lower()
    found, missing = [], []
    for section, patterns in SECTION_PATTERNS.items():
        if any(re.search(p, text_lower) for p in patterns):
            found.append(section)
        else:
            missing.append(section)
    return found, missing

# ── Contact check ─────────────────────────────────────────────────────────────
def check_contact(resume_text):
    return {
        "Email": bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume_text)),
        "Phone": bool(re.search(r'[\+\d][\d\s\-\(\)]{8,}', resume_text)),
        "LinkedIn": bool(re.search(r'linkedin', resume_text, re.I)),
        "GitHub": bool(re.search(r'github', resume_text, re.I)),
    }

# ── Format warnings ───────────────────────────────────────────────────────────
def check_format(file):
    warnings = []
    name = file.filename.lower()
    file.seek(0)
    if name.endswith(".docx"):
        try:
            doc = Document(file)
            if doc.tables:
                warnings.append("Tables detected — many ATS systems cannot parse tables correctly")
            # Check for text boxes via XML
            xml = doc.element.xml
            if 'txbxContent' in xml:
                warnings.append("Text boxes detected — ATS cannot read content inside text boxes")
            if 'graphicData' in xml or 'drawing' in xml.lower():
                warnings.append("Images or graphics detected — remove them for better ATS parsing")
        except:
            pass
    file.seek(0)
    return warnings

# ── Experience years extraction ───────────────────────────────────────────────
def extract_years_required(jd_text):
    patterns = [
        r'(\d+)\+?\s*(?:to\s*\d+)?\s*years?\s*(?:of\s*)?(?:relevant\s*)?experience',
        r'experience\s*(?:of\s*)?(\d+)\+?\s*years?',
        r'minimum\s*(\d+)\s*years?',
        r'at\s*least\s*(\d+)\s*years?',
    ]
    for p in patterns:
        m = re.search(p, jd_text, re.I)
        if m:
            return int(m.group(1))
    return None

def extract_years_resume(resume_text):
    # Count internship/experience duration mentions
    patterns = [r'(\d+)\s*years?\s*(?:of\s*)?experience', r'(\d{4})\s*[-–]\s*(\d{4}|present|current)']
    total = 0
    for m in re.finditer(patterns[1], resume_text, re.I):
        start = int(m.group(1))
        end_str = m.group(2).lower()
        end = 2025 if end_str in ['present', 'current'] else int(end_str)
        total += max(0, end - start)
    return min(total, 20)

# ── Seniority signals ─────────────────────────────────────────────────────────
SENIOR_SIGNALS = ["led","managed","owned","spearheaded","architected","designed","built","developed","launched","delivered","drove","increased","reduced","improved","optimized","automated","created","established"]
JUNIOR_SIGNALS = ["assisted","supported","helped","learning","familiar","exposure","basic","beginner","internship","trainee"]

def check_seniority(resume_text, jd_text):
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    senior_count = sum(1 for s in SENIOR_SIGNALS if s in resume_lower)
    junior_count = sum(1 for s in JUNIOR_SIGNALS if s in resume_lower)
    jd_wants_senior = any(w in jd_lower for w in ["senior","lead","manager","head","principal","architect"])
    jd_wants_junior = any(w in jd_lower for w in ["fresher","entry","graduate","junior","trainee","0-2","0-3"])
    return {
        "senior_signals": senior_count,
        "junior_signals": junior_count,
        "jd_level": "senior" if jd_wants_senior else ("junior/fresher" if jd_wants_junior else "mid-level")
    }

# ── Keyword density ───────────────────────────────────────────────────────────
def check_keyword_density(resume_text, matched_skills):
    word_count = len(resume_text.split())
    if word_count == 0:
        return "N/A", "neutral"
    density = (len(matched_skills) / max(word_count, 1)) * 100
    if density < 1:
        return f"{density:.1f}% — too low, add more relevant skills", "warning"
    elif density > 6:
        return f"{density:.1f}% — may look like keyword stuffing", "warning"
    else:
        return f"{density:.1f}% — good range", "success"

# ── Title match ───────────────────────────────────────────────────────────────
def check_title_match(resume_text, jd_text):
    titles = ["data analyst","business analyst","mis analyst","reporting analyst","operations analyst",
              "data scientist","data engineer","product analyst","marketing analyst","financial analyst",
              "hr analyst","risk analyst","senior analyst","junior analyst"]
    jd_lower = jd_text.lower()
    resume_lower = resume_text.lower()
    jd_title = next((t for t in titles if t in jd_lower), None)
    if not jd_title:
        return None, None
    match = jd_title in resume_lower
    return jd_title.title(), match

# ── Master score calculation ──────────────────────────────────────────────────
def calculate_score(matched_skills, total_skills, matched_words, total_words,
                    sections_found, contact_checks, format_warnings,
                    years_required, years_resume, title_match):

    # 1. Skill keyword score (40 pts)
    skill_score = (len(matched_skills) / max(total_skills, 1)) * 40

    # 2. General keyword score (15 pts)
    word_score = (len(matched_words) / max(total_words, 1)) * 15

    # 3. Section score (15 pts)
    section_score = (len(sections_found) / len(SECTION_PATTERNS)) * 15

    # 4. Contact score (10 pts)
    contact_score = (sum(contact_checks.values()) / len(contact_checks)) * 10

    # 5. Format score (10 pts)
    format_score = max(0, 10 - len(format_warnings) * 4)

    # 6. Experience match (5 pts)
    exp_score = 5
    if years_required:
        if years_resume >= years_required:
            exp_score = 5
        elif years_resume >= years_required - 1:
            exp_score = 3
        else:
            exp_score = 1

    # 7. Title match (5 pts)
    title_score = 5 if title_match else 2

    total = skill_score + word_score + section_score + contact_score + format_score + exp_score + title_score
    return min(int(total), 100)

def score_label(score):
    if score >= 80: return "Strong Match", "success"
    elif score >= 65: return "Good Match", "info"
    elif score >= 45: return "Moderate Match", "warning"
    return "Weak Match", "danger"

# ── Actionable tips ───────────────────────────────────────────────────────────
def generate_tips(missing_skills, sections_missing, contact_checks, format_warnings,
                  years_required, years_resume, seniority, title_name, title_match):
    tips = []

    if missing_skills:
        top = list(missing_skills.keys())[:5]
        tips.append({"icon": "🔑", "priority": "High",
                     "text": f"Add these missing skills to your resume: {', '.join(top)}. Include them naturally in your experience bullets and skills section."})

    if title_name and not title_match:
        tips.append({"icon": "🎯", "priority": "High",
                     "text": f"The JD is for a '{title_name}' but this exact title isn't visible in your resume. Add it to your summary or experience section."})

    if format_warnings:
        for w in format_warnings:
            tips.append({"icon": "⚠️", "priority": "High", "text": w + ". Fix this — ATS systems often fail to parse these elements."})

    if "Work Experience" in sections_missing:
        tips.append({"icon": "📋", "priority": "High", "text": "Work Experience section not detected. Make sure the heading is clearly labeled."})

    if "Skills" in sections_missing:
        tips.append({"icon": "📋", "priority": "High", "text": "Skills section not detected. Add a dedicated Skills or Core Skills section."})

    if not contact_checks["LinkedIn"]:
        tips.append({"icon": "🔗", "priority": "Medium", "text": "LinkedIn URL missing. Many ATS systems and recruiters specifically look for it."})

    if years_required and years_resume < years_required:
        tips.append({"icon": "📅", "priority": "Medium",
                     "text": f"JD requires {years_required}+ years of experience. Your resume shows ~{years_resume} years. Highlight all internships, projects, and freelance work."})

    if seniority["jd_level"] == "senior" and seniority["senior_signals"] < 3:
        tips.append({"icon": "💪", "priority": "Medium",
                     "text": "This is a senior role but your resume has few action verbs like 'led', 'built', 'drove', 'delivered'. Add stronger impact language."})

    if seniority["junior_signals"] > 2:
        tips.append({"icon": "✍️", "priority": "Medium",
                     "text": "Words like 'assisted', 'helped', 'familiar' weaken your resume. Replace them with stronger action verbs."})

    tips.append({"icon": "📄", "priority": "Low",
                 "text": "Always save and submit your resume as a simple single-column DOCX or PDF. No tables, no text boxes, no columns."})

    return tips

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/scan", methods=["POST"])
def scan():
    resume_file = request.files.get("resume")
    jd_text = request.form.get("jd_text", "").strip()
    if not resume_file or not jd_text:
        return render_template("index.html", error="Please upload a resume and paste the job description.")

    resume_text = extract_text(resume_file)
    if not resume_text.strip():
        return render_template("index.html", error="Could not read resume. Make sure it is a text-based PDF or DOCX.")

    # Run all checks
    format_warnings = check_format(resume_file)
    skill_matches_jd, top_words = extract_jd_keywords(jd_text)
    matched_skills, missing_skills = match_resume_skills(resume_text, skill_matches_jd)
    matched_words, missing_words = match_general_keywords(resume_text, top_words)
    sections_found, sections_missing = detect_sections(resume_text)
    contact_checks = check_contact(resume_text)
    years_required = extract_years_required(jd_text)
    years_resume = extract_years_resume(resume_text)
    seniority = check_seniority(resume_text, jd_text)
    title_name, title_match = check_title_match(resume_text, jd_text)
    density_text, density_color = check_keyword_density(resume_text, matched_skills)

    score = calculate_score(
        matched_skills, len(skill_matches_jd),
        matched_words, len(top_words),
        sections_found, contact_checks, format_warnings,
        years_required, years_resume, title_match
    )
    label, label_color = score_label(score)

    tips = generate_tips(missing_skills, sections_missing, contact_checks, format_warnings,
                         years_required, years_resume, seniority, title_name, title_match)

    # Group missing skills by category
    missing_by_cat = {}
    for canonical, info in missing_skills.items():
        cat = info["category"]
        if cat not in missing_by_cat:
            missing_by_cat[cat] = []
        missing_by_cat[cat].append(canonical)

    matched_by_cat = {}
    for canonical, info in matched_skills.items():
        cat = info["category"]
        if cat not in matched_by_cat:
            matched_by_cat[cat] = []
        matched_by_cat[cat].append(canonical)

    return render_template("result.html",
        score=score, label=label, label_color=label_color,
        matched_skills=matched_skills, missing_skills=missing_skills,
        matched_by_cat=matched_by_cat, missing_by_cat=missing_by_cat,
        matched_words=matched_words, missing_words=missing_words[:8],
        sections_found=sections_found, sections_missing=sections_missing,
        contact_checks=contact_checks, format_warnings=format_warnings,
        years_required=years_required, years_resume=years_resume,
        seniority=seniority, title_name=title_name, title_match=title_match,
        density_text=density_text, density_color=density_color,
        tips=tips,
        total_skills=len(skill_matches_jd),
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000, debug=True)
